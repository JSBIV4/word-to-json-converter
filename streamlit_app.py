import streamlit as st
import os
import json
import zipfile
from docx import Document
from datetime import datetime
import tempfile
from io import BytesIO

class WordToJsonConverter:
    def __init__(self):
        self.supported_formats = ['.docx']
    
    def validate_docx_file(self, doc_file):
        """Validate if the file is a proper DOCX file."""
        try:
            # Reset file pointer to beginning
            doc_file.seek(0)
            
            # Check if it's a valid ZIP file (DOCX files are ZIP archives)
            import zipfile
            with zipfile.ZipFile(doc_file, 'r') as zip_ref:
                # Check for required DOCX files
                required_files = ['word/document.xml', '[Content_Types].xml']
                zip_contents = zip_ref.namelist()
                
                for required_file in required_files:
                    if required_file not in zip_contents:
                        return False, f"Missing required file: {required_file}"
                
            # Reset file pointer again for Document() to use
            doc_file.seek(0)
            return True, "Valid DOCX file"
            
        except zipfile.BadZipFile:
            return False, "File is not a valid ZIP/DOCX file"
        except Exception as e:
            return False, f"File validation error: {str(e)}"

    def parse_doc_to_dict(self, doc_file):
        """Convert DOCX paragraphs with 'key: value' format into a dictionary."""
        try:
            # First validate the file
            is_valid, validation_message = self.validate_docx_file(doc_file)
            
            if not is_valid:
                st.error(f"{getattr(doc_file, 'name', 'File')}: {validation_message}")
                return None
            
            # Process the document
            doc = Document(doc_file)
            data = {
                "metadata": {
                    "source_file": getattr(doc_file, 'name', 'uploaded_file.docx'),
                    "converted_at": datetime.now().isoformat(),
                    "total_paragraphs": len(doc.paragraphs)
                },
                "content": {}
            }
            
            # Parse paragraphs
            paragraph_count = 0
            for para in doc.paragraphs:
                line = para.text.strip()
                if not line:
                    continue
                    
                paragraph_count += 1
                
                if ":" in line:
                    key, value = line.split(":", 1)
                    data["content"][key.strip()] = value.strip()
                else:
                    # Handle lines without colons as free text
                    if "free_text" not in data["content"]:
                        data["content"]["free_text"] = []
                    data["content"]["free_text"].append(line)
            
            # Update actual paragraph count with content
            data["metadata"]["content_paragraphs"] = paragraph_count
            
            if not data["content"]:
                st.warning(f"{getattr(doc_file, 'name', 'File')}: No content found (document might be empty or use unsupported formatting)")
            
            return data
            
        except Exception as e:
            st.error(f"Error processing {getattr(doc_file, 'name', 'document')}: {str(e)}")
            return None

def main():
    st.set_page_config(
        page_title="Word to JSON Converter",
        page_icon="📄",
        layout="wide"
    )
    
    st.title("📄 Word Document to JSON Converter")
    st.markdown("Upload your Word documents (.docx) and convert them to JSON format")
    
    converter = WordToJsonConverter()
    
    # Sidebar with instructions
    with st.sidebar:
        st.header("Instructions")
        st.markdown("""
        1. Upload one or multiple .docx files
        2. The converter looks for 'key: value' pairs in your document
        3. Download the converted JSON files
        
        **Supported format:**
        ```
        Name: John Doe
        Age: 30
        City: New York
        ```
        
        **Output JSON:**
        ```json
        {
          "content": {
            "Name": "John Doe",
            "Age": "30", 
            "City": "New York"
          }
        }
        ```
        """)
    
    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("Upload Documents")
        uploaded_files = st.file_uploader(
            "Choose Word documents",
            type=['docx'],
            accept_multiple_files=True,
            help="Upload one or more .docx files to convert to JSON"
        )
        
        if uploaded_files:
            st.success(f"Uploaded {len(uploaded_files)} file(s)")
            
            # Show file details
            for file in uploaded_files:
                st.write(f"{file.name} ({file.size:,} bytes)")
    
    with col2:
        st.header("Conversion Options")
        include_metadata = st.checkbox("Include metadata", value=True)
        pretty_format = st.checkbox("Pretty format JSON", value=True)
        
    if uploaded_files:
        st.header("Convert Documents")
        
        if st.button("Convert All Documents", type="primary"):
            converted_files = {}
            failed_files = []
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for i, uploaded_file in enumerate(uploaded_files):
                status_text.text(f"Converting {uploaded_file.name}...")
                
                # Show file info for debugging
                st.write(f"🔍 Processing: {uploaded_file.name} ({uploaded_file.size:,} bytes)")
                
                # Convert document
                data = converter.parse_doc_to_dict(uploaded_file)
                
                if data:
                    if not include_metadata:
                        data = data.get("content", {})
                    
                    # Format JSON
                    if pretty_format:
                        json_str = json.dumps(data, indent=2, ensure_ascii=False)
                    else:
                        json_str = json.dumps(data, ensure_ascii=False)
                    
                    # Store converted data
                    json_filename = uploaded_file.name.replace('.docx', '.json')
                    converted_files[json_filename] = json_str
                    st.success(f"Successfully converted {uploaded_file.name}")
                else:
                    failed_files.append(uploaded_file.name)
                    st.error(f"Failed to convert {uploaded_file.name}")
                    
                progress_bar.progress((i + 1) / len(uploaded_files))
            
            # Summary
            if converted_files:
                status_text.text(f"Conversion complete! {len(converted_files)} successful, {len(failed_files)} failed")
            else:
                status_text.text("No files were successfully converted")
                
            if failed_files:
                st.error(f"Failed to convert: {', '.join(failed_files)}")
                st.info("**Troubleshooting tips:**\n- Make sure files are valid .docx format (not .doc)\n- Try re-saving the file in Microsoft Word\n- Check if file is corrupted")
            
            # Display results
            st.header("Download Results")
            
            if len(converted_files) == 1:
                # Single file download
                filename, content = list(converted_files.items())[0]
                st.download_button(
                    label=f"Download {filename}",
                    data=content,
                    file_name=filename,
                    mime="application/json"
                )
            else:
                # Multiple files - create ZIP
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    for filename, content in converted_files.items():
                        zip_file.writestr(filename, content)
                
                zip_buffer.seek(0)
                
                st.download_button(
                    label=f"Download All Files ({len(converted_files)} files)",
                    data=zip_buffer.getvalue(),
                    file_name=f"converted_documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                    mime="application/zip"
                )
            
            # Preview section
            st.header("Preview Results")
            preview_file = st.selectbox("Select file to preview:", list(converted_files.keys()))
            
            if preview_file:
                st.code(converted_files[preview_file], language='json')
                
                # Show statistics
                data = json.loads(converted_files[preview_file])
                if isinstance(data, dict):
                    content_keys = data.get("content", data) if "content" in data else data
                    if isinstance(content_keys, dict):
                        st.info(f"Found {len(content_keys)} key-value pairs in {preview_file}")

if __name__ == "__main__":
    main()