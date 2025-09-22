import os
import json
import zipfile
from docx import Document
from pathlib import Path
from datetime import datetime
import tempfile

class WordToJsonConverter:
    def __init__(self):
        self.supported_formats = ['.docx']
    
    def parse_doc_to_dict(self, doc_path):
        """Convert DOCX paragraphs with 'key: value' format into a dictionary."""
        
        # Skip Word lock/temp files
        if os.path.basename(doc_path).startswith("~$"):
            print(f"Skipping lock file: {doc_path}")
            return None
        
        try:
            doc = Document(doc_path)
            data = {
                "metadata": {
                    "source_file": os.path.basename(doc_path),
                    "converted_at": datetime.now().isoformat(),
                    "total_paragraphs": len(doc.paragraphs)
                },
                "content": {}
            }
            
            # Parse paragraphs
            for para in doc.paragraphs:
                line = para.text.strip()
                if not line:
                    continue
                if ":" in line:
                    key, value = line.split(":", 1)
                    data["content"][key.strip()] = value.strip()
                else:
                    # Handle lines without colons as free text
                    if "free_text" not in data["content"]:
                        data["content"]["free_text"] = []
                    data["content"]["free_text"].append(line)
            
            return data
            
        except Exception as e:
            print(f"Error processing {doc_path}: {str(e)}")
            return None
    
    def convert_single_file(self, input_path, output_path=None):
        """Convert a single Word document to JSON."""
        if not output_path:
            output_path = input_path.replace('.docx', '.json')
        
        data = self.parse_doc_to_dict(input_path)
        if data:
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            return output_path
        return None
    
    def convert_folder(self, docx_folder, json_folder=None):
        """Convert all Word documents in a folder (recursively)."""
        if not json_folder:
            json_folder = docx_folder
        
        os.makedirs(json_folder, exist_ok=True)
        converted_files = []
        
        # Walk through DOCX_FOLDER recursively
        for root, dirs, files in os.walk(docx_folder):
            for filename in files:
                if filename.endswith(".docx"):
                    doc_path = os.path.join(root, filename)
                    
                    # Parse DOCX to dictionary
                    data = self.parse_doc_to_dict(doc_path)
                    if not data:
                        continue
                    
                    # Determine relative path to preserve folder structure
                    rel_path = os.path.relpath(root, docx_folder)
                    target_folder = os.path.join(json_folder, rel_path)
                    os.makedirs(target_folder, exist_ok=True)
                    
                    # Save JSON
                    base_name = os.path.splitext(filename)[0]
                    json_path = os.path.join(target_folder, f"{base_name}.json")
                    
                    with open(json_path, "w", encoding="utf-8") as f:
                        json.dump(data, f, indent=2, ensure_ascii=False)
                    
                    converted_files.append(json_path)
                    print(f"[✓] Converted {doc_path} → {json_path}")
        
        return converted_files

# Command line usage example
if __name__ == "__main__":
    converter = WordToJsonConverter()
    
    # Example usage - update these paths as needed
    DOCX_FOLDER = "sample_data/7. Critical Illness_all"
    JSON_FOLDER = "sample_data/7. Critical Illness_all"
    
    if os.path.exists(DOCX_FOLDER):
        converted = converter.convert_folder(DOCX_FOLDER, JSON_FOLDER)
        print(f"\nConversion complete! Converted {len(converted)} files.")
    else:
        print(f"Folder not found: {DOCX_FOLDER}")
        print("Please update the DOCX_FOLDER path in the script.")